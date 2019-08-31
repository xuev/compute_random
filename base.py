#!/usr/bin/env python
# coding: utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from random import randrange
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import time


# 加载配置文件信息
def load_conf(config_name):
    config_dic = {}
    with open(config_name, 'r') as f:
        lines = f.readlines()
        for config in lines:
            config_str = config.strip('\n')
            config_key = config_str.split('=')[0]
            config_value = config_str.split('=')[1]
            config_dic[config_key] = config_value
    return config_dic


# 生成减法列表
def get_sub_list():
    sub_list = []
    for x in range(1, int(configs_dic['range'])):
        for y in range(1, x + 1):
            if x != y:
                sub_item = str(x) + " - " + str(y) + " = "
                sub_list.append(sub_item)
    return sub_list


# 生成加法列表
def get_add_list():
    add_list = []
    for a in range(1, int(configs_dic['range'])):
        for b in range(1, int(configs_dic['range'])):
            add_item = str(a) + " + " + str(b) + " = "
            add_list.append(add_item)
    return add_list


# 生成目标列表
def get_target_list(add_li, sub_li):
    all = add_li + sub_li
    target_list = []
    for i in range(int(configs_dic['page'])):
        random_index = randrange(0, len(all))
        target_list.append(all[random_index])
    return target_list


# 把目标列表写入word
def write_to_word(target_list):
    row_size = int(len(target_list) // 3)
    left_target_list = target_list[:row_size]
    middle_target_list = target_list[row_size:row_size * 2]
    right_target_list = target_list[row_size * 2:]
    # 打开文档
    document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))

    # 加入标题
    run = document.add_heading('', level=0).add_run(configs_dic['range'] + u'以内加减法', 0)
    run.font.name = u'微软雅黑'
    run.font.size = Pt(20)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # document.add_heading(u'测试10以内加减，加强练习', 0)

    # 添加文本

    for i in range(row_size):
        paragraph = document.add_paragraph()
        # 设置中文字体
        run = paragraph.add_run(
            left_target_list[i] + "\t    " + middle_target_list[i] + "\t    " + right_target_list[i])
        run.font.name = u'宋体'
        run.font.size = Pt(20)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 保存文件
    document.save(configs_dic['range'] + u'以内加减法 - ' + time.strftime("%Y%m%d_%H%M%S", time.localtime()) + '.docx')


if __name__ == '__main__':
    # 读取配置文件信息,range=10 范围，page=45,每页算术个数
    config_name = './conf/conf.ini'
    configs_dic = load_conf(config_name)
    # 生成加、减法列表
    add_li = get_add_list()
    sub_li = get_sub_list()
    # 生成目标列表，45个
    target_list = get_target_list(add_li, sub_li)
    # 写入word
    write_to_word(target_list)
